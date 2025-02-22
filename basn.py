#pip install pymupdf python-docx
import pymupdf as fitz
import docx
from docx.shared import Inches

def extract_text_from_pdf(pdf_path, word_path):

    try:
        pdf_document = fitz.open(pdf_path)
        doc = docx.Document()

        for page_num in range(pdf_document.page_count):
            page = pdf_document.load_page(page_num)

            text = page.get_text()
            if text:
                doc.add_paragraph(text)
        image_list = page.get_images()

        for img_ind, img in enumerate(image_list):
            xref = img[0]
            base_image = pdf_document.extract_image(xref)
            image_bytes  = base_image["image"]
            image_ext = base_image["ext"]

            image_path_temp = f"image_{img_ind}.{image_ext}"
            with open(image_path_temp,"wb") as f:
                f.write(image_bytes)

            doc.add_picture(image_path_temp, width=Inches(3), height=Inches(3))

            import os
            os.remove(image_path_temp)

        if page_num < pdf_document.page_count -1:
            doc.add_page_break()

        doc.save(word_path)
        print("Text and images extracted and saved to Word document sucessfully.")

    except Exception as e:
         print(f"error: {e}")

pdf_file = "scanDoc/wirausaha.pdf"
word_file = "scanDoc/suart.docx"

extract_text_from_pdf("wirausaha.pdf", "suart.docx")
